#!/usr/bin/env python3
"""Convierte el instrumento-validacion-docente.md a .docx con marca UA básica.

Uso: python3 build_instrumento_docx.py <instrumento.md> [salida.docx]

Requiere python-docx (pip install python-docx). Si no está instalado, informa y sale sin error fatal.
Conversión markdown->docx intencionalmente simple (encabezados, listas, tablas, párrafos y casillas);
no pretende fidelidad tipográfica total: el objetivo es entregar un editable para el docente.
"""
import sys
import os
import re

def main():
    if len(sys.argv) < 2:
        print("Uso: python3 build_instrumento_docx.py <instrumento.md> [salida.docx]", file=sys.stderr)
        return 2
    md_path = sys.argv[1]
    out_path = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(md_path)[0] + ".docx"
    if not os.path.exists(md_path):
        print(f"No existe: {md_path}", file=sys.stderr)
        return 2
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("python-docx no está instalado. Instala con: pip install python-docx", file=sys.stderr)
        print("El instrumento .md sigue siendo válido para socializar con el docente.", file=sys.stderr)
        return 0

    UA_GREEN = RGBColor(0x1A, 0x24, 0x03)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        # Tablas markdown: bloque de líneas que empiezan con '|'.
        if line.strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in block]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]  # quitar separador ---
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                for ri, r in enumerate(rows):
                    for ci, c in enumerate(r):
                        table.cell(ri, ci).text = re.sub(r"[*_`]", "", c)
            continue
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs: run.font.color.rgb = UA_GREEN
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs: run.font.color.rgb = UA_GREEN
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip().startswith(("- ", "* ")):
            doc.add_paragraph(re.sub(r"[*_`]", "", line.strip()[2:]), style="List Bullet")
        elif line.strip().startswith("> "):
            p = doc.add_paragraph(re.sub(r"[*_`]", "", line.strip()[2:]))
            p.paragraph_format.left_indent = Pt(18)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip():
            doc.add_paragraph(re.sub(r"[*_`]", "", line))
        i += 1

    doc.save(out_path)
    print(f"Escrito: {out_path}")
    return 0

if __name__ == "__main__":
    sys.exit(main())
