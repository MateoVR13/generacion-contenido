from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKDIR = Path(__file__).resolve().parent
SOURCE_MD = WORKDIR / "ds_9_ensamble_final_documento_saberes_uoua_emprendimiento.md"
OUTPUT_DOCX = WORKDIR / "documento_saberes_emprendimiento_uoua_final_uamerica.docx"
BRAND_DIR = Path(r"C:\Users\Usuario\.codex\skills\manual-de-marca-u-america")
LOGO_UAMERICA = BRAND_DIR / "assets" / "logo-uamerica-1-linea-full-color.png"
LOGO_CEIS = BRAND_DIR / "assets" / "logo-ceis-uamerica-full-color.png"


GREEN_DARK = RGBColor(0x1A, 0x24, 0x03)
LIME = RGBColor(0xC8, 0xFF, 0x01)
YELLOW = RGBColor(0xFF, 0xCE, 0x00)
GRAY_DARK = RGBColor(0x3D, 0x3D, 0x3F)
GRAY_TEXT = RGBColor(0x60, 0x60, 0x60)
GRAY_LINE = "D9DED1"
LIGHT_FILL = "F7F9F2"
FONT = "Montserrat"
FALLBACK_FONT = "Arial"


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.18):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="A7A9AA", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def mark_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_picture_alt(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total = sum(widths_dxa)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_bottom_border(paragraph, color="C8FF01", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_markdown_runs(paragraph, text, default_size=10.5, default_color=GRAY_DARK):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=max(default_size - 0.5, 8), color=GRAY_DARK)
        else:
            clean = part.replace("**", "")
            run = paragraph.add_run(clean)
            set_run_font(run, size=default_size, color=default_color)


def clean_inline(text):
    return text.strip().replace("**", "").replace("`", "")


def split_table_row(line):
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [c.strip() for c in stripped.split("|")]


def is_separator_row(line):
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells)


def compute_widths(rows, total=9360):
    col_count = max(len(row) for row in rows)
    weights = [1.0] * col_count
    for row in rows[:8]:
        for i, cell in enumerate(row):
            length = len(clean_inline(cell))
            weights[i] = max(weights[i], min(5.0, 0.6 + length / 28))
    min_width = 850 if col_count >= 5 else 1100
    raw = [max(min_width, int(total * w / sum(weights))) for w in weights]
    diff = total - sum(raw)
    raw[-1] += diff
    if raw[-1] < min_width:
        deficit = min_width - raw[-1]
        raw[-1] = min_width
        for i in range(len(raw) - 1):
            take = min(deficit, max(0, raw[i] - min_width))
            raw[i] -= take
            deficit -= take
            if deficit <= 0:
                break
    return raw


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = GRAY_DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Title", 24, GREEN_DARK, 0, 8),
        ("Heading 1", 16, GREEN_DARK, 18, 8),
        ("Heading 2", 13, GREEN_DARK, 14, 6),
        ("Heading 3", 11.5, GRAY_DARK, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True if name != "Title" else False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18


def setup_sections(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Documento de Saberes | CEIS UAmerica")
    set_run_font(run, size=8.5, color=GRAY_TEXT, bold=True)
    add_bottom_border(p, color="D9DED1", size="4")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(fp.add_run("Universidad de America - CEIS | Vigilada MinEducacion | Pagina "), size=8, color=GRAY_TEXT)
    add_page_number(fp)

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(ffp.add_run("Codigo SNIES 1715 | Vigilada MinEducacion"), size=8.5, color=GRAY_TEXT)


def add_cover(doc):
    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=0)

    if LOGO_UAMERICA.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(str(LOGO_UAMERICA), width=Inches(4.4))
        set_picture_alt(
            picture,
            "Logo Universidad de America",
            "Logo institucional horizontal de la Universidad de America.",
        )
        set_paragraph_spacing(p, after=14)
    if LOGO_CEIS.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(str(LOGO_CEIS), width=Inches(1.45))
        set_picture_alt(
            picture,
            "Logo CEIS UAmerica",
            "Logo del Centro de Emprendimiento e Innovacion Sostenible CEIS UAmerica.",
        )
        set_paragraph_spacing(p, after=26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento de Saberes")
    set_run_font(r, size=14, color=GREEN_DARK, bold=True)
    set_paragraph_spacing(p, after=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Validacion temprana de oportunidades de emprendimiento")
    set_run_font(r, size=23, color=GREEN_DARK, bold=True)
    set_paragraph_spacing(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("De problema real a decision Gate 1")
    set_run_font(r, size=14, color=GRAY_DARK)
    set_paragraph_spacing(p, after=18)
    add_bottom_border(p, color="C8FF01", size="16")

    meta = [
        ("Asignatura", "Emprendimiento 1 / Fundamentos de Emprendimiento"),
        ("Dependencia", "Centro de Emprendimiento e Innovacion Sostenible - CEIS"),
        ("Modalidad", "Virtual"),
        ("Creditos", "3"),
        ("Version", "Ensamble final DS-9"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_borders(table, color="D9DED1", size="4")
    widths = [1900, 7460]
    for label, value in meta:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], "F0F5E8")
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, top=90, bottom=90)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=0, line=1.1)
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, color=GREEN_DARK if idx == 0 else GRAY_DARK, bold=(idx == 0))
    set_table_geometry(table, widths)
    mark_row_as_header(table.rows[0])

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=18, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Universidad de America | CEIS UAmerica"), size=10, color=GRAY_TEXT, bold=True)
    doc.add_page_break()


def extract_body_markdown():
    text = SOURCE_MD.read_text(encoding="utf-8")
    start = text.find("# Documento de Saberes integrado")
    if start == -1:
        start = 0
    return text[start:].strip()


def collect_headings(markdown):
    headings = []
    for line in markdown.splitlines():
        m = re.match(r"^(#{1,3})\s+(.+)$", line.strip())
        if m:
            title = clean_inline(m.group(2))
            if "Subfase DS-9" not in title:
                headings.append((len(m.group(1)), title))
    return headings


def add_index(doc, headings):
    p = doc.add_paragraph()
    r = p.add_run("Indice de contenidos")
    set_run_font(r, size=16, color=GREEN_DARK, bold=True)
    set_paragraph_spacing(p, before=0, after=8)

    for level, title in headings:
        if level > 2:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2, line=1.1)
        p.paragraph_format.left_indent = Inches(0.0 if level == 1 else 0.25)
        r = p.add_run(title)
        set_run_font(r, size=10.2 if level == 1 else 9.6, color=GREEN_DARK if level == 1 else GRAY_DARK, bold=(level == 1))
    doc.add_page_break()


def add_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.0)
        p.paragraph_format.left_indent = Inches(0.18)
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.8, color=GRAY_DARK)


def add_markdown_table(doc, lines):
    rows = [split_table_row(line) for line in lines if line.strip()]
    has_header = len(rows) >= 2 and is_separator_row(lines[1])
    if has_header:
        header = rows[0]
        body = rows[2:]
    else:
        header = None
        body = rows
    data = ([header] if header else []) + body
    if not data:
        return
    col_count = max(len(row) for row in data)
    for row in data:
        row.extend([""] * (col_count - len(row)))
    widths = compute_widths(data)
    table = doc.add_table(rows=len(data), cols=col_count)
    table.style = "Table Grid"
    set_table_borders(table, color="A7A9AA", size="4")
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            if has_header and r_idx == 0:
                set_cell_shading(cell, "1A2403")
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "FFFFFF")
            else:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.05)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            size = 8.3 if col_count >= 5 else 8.8
            color = RGBColor(255, 255, 255) if has_header and r_idx == 0 else GRAY_DARK
            add_markdown_runs(p, clean_inline(value), default_size=size, default_color=color)
            for run in p.runs:
                run.bold = True if has_header and r_idx == 0 else run.bold
    if table.rows:
        mark_row_as_header(table.rows[0])
    after = doc.add_paragraph()
    set_paragraph_spacing(after, after=2)


def add_body_from_markdown(doc, markdown):
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            p = doc.add_paragraph()
            add_bottom_border(p, color="D9DED1", size="6")
            set_paragraph_spacing(p, before=2, after=8)
            i += 1
            continue

        if "|" in stripped and stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            p = doc.add_paragraph(style=style)
            add_markdown_runs(p, text, default_size=16 if level == 1 else 13 if level == 2 else 11.5, default_color=GREEN_DARK if level <= 2 else GRAY_DARK)
            if level == 1:
                add_bottom_border(p, color="C8FF01", size="8")
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_markdown_runs(p, numbered.group(1), default_size=10.2, default_color=GRAY_DARK)
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_runs(p, bullet.group(1), default_size=10.2, default_color=GRAY_DARK)
            i += 1
            continue

        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=7, line=1.2)
        add_markdown_runs(p, stripped, default_size=10.5, default_color=GRAY_DARK)
        i += 1


def main():
    markdown = extract_body_markdown()
    doc = Document()
    configure_styles(doc)
    setup_sections(doc)
    add_cover(doc)
    add_index(doc, collect_headings(markdown))
    add_body_from_markdown(doc, markdown)

    core = doc.core_properties
    core.title = "Documento de Saberes - Emprendimiento 1"
    core.subject = "Validacion temprana de oportunidades de emprendimiento"
    core.author = "Universidad de America - CEIS"
    core.comments = "Generado desde DS-9 con lineamientos de marca UAmerica/CEIS."

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
