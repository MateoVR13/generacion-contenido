from __future__ import annotations

import json
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


DOCX = Path("documento_saberes_final_45_paginas_emprendimiento_uamerica.docx")
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def main():
    doc = Document(DOCX)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    headings = []
    for paragraph in doc.paragraphs:
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading") and paragraph.text.strip():
            headings.append({"style": style, "text": paragraph.text.strip()})

    with zipfile.ZipFile(DOCX) as zf:
        names = zf.namelist()
        media = [n for n in names if n.startswith("word/media/")]
        document_xml = zf.read("word/document.xml")
        root = ET.fromstring(document_xml)
        sect_pr = root.findall(".//w:sectPr", NS)
        tbls = root.findall(".//w:tbl", NS)
        body_text = "\n".join(paragraphs)

    report = {
        "file": str(DOCX.resolve()),
        "exists": DOCX.exists(),
        "size_bytes": DOCX.stat().st_size,
        "paragraphs_nonempty": len(paragraphs),
        "tables": len(doc.tables),
        "tables_xml": len(tbls),
        "sections": len(doc.sections),
        "section_properties_xml": len(sect_pr),
        "inline_shapes": len(doc.inline_shapes),
        "media_files": len(media),
        "heading_count": len(headings),
        "first_headings": headings[:20],
        "has_presentacion": "Presentacion del documento" in body_text or "Presentación del documento" in body_text,
        "has_marco": "Marco conceptual" in body_text,
        "has_momento_1": "Momento 1" in body_text,
        "has_momento_2": "Momento 2" in body_text,
        "has_momento_3": "Momento 3" in body_text,
        "has_momento_4": "Momento 4" in body_text,
        "has_glosario": "Glosario" in body_text,
        "has_referencias": "Referencias" in body_text,
        "has_anexos": "Anexos" in body_text,
        "core_title": doc.core_properties.title,
        "core_author": doc.core_properties.author,
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
