from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

import build_ds9_word_uamerica as base


WORKDIR = Path(__file__).resolve().parent
SOURCE_MD = WORKDIR / "documento_saberes_final_45_paginas_emprendimiento.md"
OUTPUT_DOCX = WORKDIR / "documento_saberes_final_45_paginas_emprendimiento_uamerica.docx"


def extract_body_markdown():
    text = SOURCE_MD.read_text(encoding="utf-8")
    marker = "## Tabla de control editorial y proyección de páginas"
    start = text.find(marker)
    if start == -1:
        marker = "# Presentación del documento"
        start = text.find(marker)
    if start == -1:
        start = 0
    return text[start:].strip()


def add_cover(doc):
    for _ in range(2):
        p = doc.add_paragraph()
        base.set_paragraph_spacing(p, after=0)

    if base.LOGO_UAMERICA.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(str(base.LOGO_UAMERICA), width=Inches(4.35))
        base.set_picture_alt(
            picture,
            "Logo Universidad de America",
            "Logo institucional horizontal de la Universidad de America.",
        )
        base.set_paragraph_spacing(p, after=14)

    if base.LOGO_CEIS.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(str(base.LOGO_CEIS), width=Inches(1.42))
        base.set_picture_alt(
            picture,
            "Logo CEIS UAmerica",
            "Logo del Centro de Emprendimiento e Innovacion Sostenible CEIS UAmerica.",
        )
        base.set_paragraph_spacing(p, after=24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento de Saberes")
    base.set_run_font(r, size=14, color=base.GREEN_DARK, bold=True)
    base.set_paragraph_spacing(p, after=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Validacion temprana de oportunidades de emprendimiento")
    base.set_run_font(r, size=23, color=base.GREEN_DARK, bold=True)
    base.set_paragraph_spacing(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("De problema real a decision Gate 1")
    base.set_run_font(r, size=14, color=base.GRAY_DARK)
    base.set_paragraph_spacing(p, after=18)
    base.add_bottom_border(p, color="C8FF01", size="16")

    meta = [
        ("Asignatura", "Emprendimiento 1 / Fundamentos de Emprendimiento"),
        ("Dependencia", "Centro de Emprendimiento e Innovacion Sostenible - CEIS"),
        ("Modalidad", "Virtual"),
        ("Creditos", "3"),
        ("Extension", "40 a 45 paginas maquetadas"),
        ("Version", "Documento final aprobado"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    base.set_table_borders(table, color="D9DED1", size="4")
    widths = [1900, 7460]
    for label, value in meta:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        base.set_cell_shading(cells[0], "F0F5E8")
        for idx, cell in enumerate(cells):
            base.set_cell_width(cell, widths[idx])
            base.set_cell_margins(cell, top=90, bottom=90)
            for paragraph in cell.paragraphs:
                base.set_paragraph_spacing(paragraph, after=0, line=1.1)
                for run in paragraph.runs:
                    base.set_run_font(
                        run,
                        size=9.5,
                        color=base.GREEN_DARK if idx == 0 else base.GRAY_DARK,
                        bold=(idx == 0),
                    )
    base.set_table_geometry(table, widths)
    base.mark_row_as_header(table.rows[0])

    p = doc.add_paragraph()
    base.set_paragraph_spacing(p, before=18, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run_font(
        p.add_run("Universidad de America | CEIS UAmerica"),
        size=10,
        color=base.GRAY_TEXT,
        bold=True,
    )
    doc.add_page_break()


def main():
    base.SOURCE_MD = SOURCE_MD
    base.OUTPUT_DOCX = OUTPUT_DOCX
    base.extract_body_markdown = extract_body_markdown
    base.add_cover = add_cover

    markdown = extract_body_markdown()
    doc = base.Document()
    base.configure_styles(doc)
    base.setup_sections(doc)
    add_cover(doc)
    base.add_index(doc, base.collect_headings(markdown))
    base.add_body_from_markdown(doc, markdown)

    core = doc.core_properties
    core.title = "Documento de Saberes - Emprendimiento 1"
    core.subject = "Validacion temprana de oportunidades de emprendimiento"
    core.author = "Universidad de America - CEIS"
    core.comments = "Generado desde documento_saberes_final_45_paginas_emprendimiento.md con lineamientos de marca UAmerica/CEIS."

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
